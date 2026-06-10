"""Сборка ВКР.md -> ВКР.docx с оформлением по ГОСТ 7.32-2017.

TNR 14, интервал 1.5, поля Л30/П10/В20/Н20 мм, абзац 1.25 см, выравнивание по
ширине; структурные заголовки по центру прописными, разделы с новой страницы;
рисунки по центру, таблицы с рамкой, листинги Courier New, номера страниц снизу.
"""
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

# Ширина полосы набора: A4 (210 мм) минус поля 30/10 мм
TEXT_WIDTH_MM = 170

DIPLOMA_DIR = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(DIPLOMA_DIR, "ВКР.md")
OUT = os.path.join(DIPLOMA_DIR, "ВКР.docx")

STRUCT = {
    "РЕФЕРАТ", "ABSTRACT", "СОДЕРЖАНИЕ", "ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ",
    "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ", "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
}


def set_rfonts(element, name):
    rpr = element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), name)


def setup_styles(doc):
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(14)
    set_rfonts(n.element, "Times New Roman")
    pf = n.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(1.25)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for i in (1, 2, 3):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(14)
        h.font.bold = True
        h.font.color.rgb = RGBColor(0, 0, 0)
        set_rfonts(h.element, "Times New Roman")
        hp = h.paragraph_format
        hp.line_spacing = 1.5
        hp.space_before = Pt(12)
        hp.space_after = Pt(6)
        hp.keep_with_next = True


def setup_section(doc):
    s = doc.sections[0]
    s.left_margin, s.right_margin = Mm(30), Mm(10)
    s.top_margin, s.bottom_margin = Mm(20), Mm(20)
    p = s.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    for kind, txt in (("begin", None), ("instr", "PAGE"), ("end", None)):
        if kind == "instr":
            e = OxmlElement("w:instrText")
            e.set(qn("xml:space"), "preserve")
            e.text = txt
        else:
            e = OxmlElement("w:fldChar")
            e.set(qn("w:fldCharType"), kind)
        run._r.append(e)


def add_runs(p, text):
    for i, part in enumerate(text.split("**")):
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = "Times New Roman"
        r.font.size = Pt(14)
        if i % 2 == 1:
            r.bold = True


class Builder:
    def __init__(self, doc):
        self.doc = doc
        self.first_h1 = True

    def heading(self, text, level):
        text = text.replace("**", "").strip()
        p = self.doc.add_paragraph(style=f"Heading {level}")
        structural = level == 1 and (text in STRUCT or text.startswith("ПРИЛОЖЕНИЕ"))
        if level == 1:
            if not self.first_h1:
                p.paragraph_format.page_break_before = True
            self.first_h1 = False
        if structural:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
        else:
            # Методичка: заголовки разделов и подразделов — с абзацного отступа
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(1.25)
        r = p.add_run(text)
        r.bold = True

    FORMULA_RE = re.compile(r"^(.{4,}?)\s*\((\d+)\)$")

    def body(self, text):
        if text.startswith("- ") or text.startswith("* "):
            text = "– " + text[2:]
        # Формула с номером: по центру, номер по правому краю (методичка, п. 5.7)
        m = self.FORMULA_RE.match(text)
        if m and any(ch in m.group(1) for ch in "=≥≤≈"):
            self.formula(m.group(1).strip(), m.group(2))
            return
        p = self.doc.add_paragraph(style="Normal")
        if text.startswith("где "):  # расшифровка формулы — без абзацного отступа
            p.paragraph_format.first_line_indent = Cm(0)
        add_runs(p, text)

    def formula(self, body, num):
        p = self.doc.add_paragraph(style="Normal")
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.tab_stops.add_tab_stop(Mm(TEXT_WIDTH_MM / 2), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Mm(TEXT_WIDTH_MM), WD_TAB_ALIGNMENT.RIGHT)
        add_runs(p, "\t" + body + "\t(" + num + ")")

    def caption(self, text, center, size=14, before=0, after=0):
        p = self.doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        add_runs(p, text)
        for r in p.runs:
            r.font.size = Pt(size)

    def image(self, path):
        full = os.path.join(DIPLOMA_DIR, path)
        p = self.doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        if os.path.exists(full):
            p.add_run().add_picture(full, width=Mm(155))
        else:
            add_runs(p, f"[рисунок отсутствует: {path}]")

    def code(self, lines, caption):
        if caption:
            self.caption(caption, center=False)
        for ln in lines:
            p = self.doc.add_paragraph(style="Normal")
            pf = p.paragraph_format
            pf.first_line_indent = Cm(0)
            pf.left_indent = Cm(0.5)
            pf.line_spacing = 1.0
            r = p.add_run(ln if ln.strip() else " ")
            r.font.name = "Courier New"
            r.font.size = Pt(12)  # методичка, п. 5.11: кегль листинга 14/12/10
            set_rfonts(r._element, "Courier New")

    def table(self, rows):
        ncol = len(rows[0])
        t = self.doc.add_table(rows=0, cols=ncol)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ri, row in enumerate(rows):
            cells = t.add_row().cells
            for ci in range(ncol):
                val = row[ci] if ci < len(row) else ""
                cell = cells[ci]
                para = cell.paragraphs[0]
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.line_spacing = 1.0
                add_runs(para, val)
                for r in para.runs:
                    r.font.size = Pt(12)
                    if ri == 0:
                        r.bold = True
        # Отбивка таблицы от последующего текста (методичка: 12 пт)
        sep = self.doc.add_paragraph(style="Normal")
        sep.paragraph_format.line_spacing = 1.0
        sep.paragraph_format.space_after = Pt(0)
        r = sep.add_run(" ")
        r.font.size = Pt(2)


def is_table_row(line):
    return line.strip().startswith("|")


def parse(md, b):
    lines = md.split("\n")
    i, n = 0, len(lines)
    para = []
    caption_lst = None

    def flush():
        nonlocal para
        if para:
            b.body(" ".join(s.strip() for s in para))
            para = []

    while i < n:
        raw = lines[i]
        line = raw.rstrip("\n")
        s = line.strip()

        if s.startswith("```"):
            flush()
            code_lines = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            b.code(code_lines, caption_lst)
            caption_lst = None
            i += 1
            continue

        if is_table_row(line):
            flush()
            rows = []
            while i < n and is_table_row(lines[i]):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(set(c) <= set("-: ") for c in cells):  # пропустить разделитель
                    rows.append(cells)
                i += 1
            if rows:
                b.table(rows)
            continue

        if s == "" or s == "---":
            flush()
            i += 1
            continue
        if s.startswith(">") or s.startswith("# "):  # черновая пометка и титул — пропустить
            flush()
            i += 1
            continue
        if s.startswith("#### "):
            flush(); b.heading(s[5:], 3); i += 1; continue
        if s.startswith("### "):
            flush(); b.heading(s[4:], 2); i += 1; continue
        if s.startswith("## "):
            flush(); b.heading(s[3:], 1); i += 1; continue
        m = re.match(r"!\[.*?\]\((.+?)\)", s)
        if m:
            flush(); b.image(m.group(1)); i += 1; continue
        if s.startswith("Рисунок "):
            # Подпись рисунка: по центру, кегль 13, отбивка 12 пт (методичка, п. 5.5)
            flush(); b.caption(s, center=True, size=13, after=12); i += 1; continue
        if s.startswith("Таблица "):
            # Название таблицы: по центру, отбивка от текста 12 пт (методичка, п. 5.6)
            flush(); b.caption(s, center=True, before=12); i += 1; continue
        if s.startswith("Листинг "):
            flush(); caption_lst = s; i += 1; continue
        if re.match(r"^\d+\.\s", s) or re.match(r"^[-*]\s", s):
            flush(); para = [s]; i += 1; continue
        if raw[:1] in (" ", "\t") and para:  # продолжение абзаца/пункта
            para.append(s); i += 1; continue
        para.append(s)
        i += 1
    flush()


def main():
    with open(SRC, encoding="utf-8") as f:
        md = f.read()
    doc = Document()
    setup_styles(doc)
    setup_section(doc)
    b = Builder(doc)
    parse(md, b)
    # Полные листинги исходного кода в приложение В
    listings = [
        ("../ble-scanner/trajectory.py",
         "Листинг В.4 — Модуль анализа траектории (trajectory.py)"),
        ("../ble-scanner/hm10.py",
         "Листинг В.5 — Модуль управления HM10 (hm10.py)"),
        ("../hm10_lock/lib/hm10_service.dart",
         "Листинг В.6 — Сервис HM10 мобильного приложения (hm10_service.dart)"),
        ("../hm10_lock/lib/virtual_lock_db.dart",
         "Листинг В.7 — Виртуальная база замков (virtual_lock_db.dart)"),
        ("../hm10_lock/lib/lock_send_page.dart",
         "Листинг В.8 — Экран управления замком (lock_send_page.dart)"),
    ]
    for rel, cap in listings:
        full = os.path.join(DIPLOMA_DIR, rel)
        if os.path.exists(full):
            doc.add_page_break()
            with open(full, encoding="utf-8") as fh:
                b.code(fh.read().split("\n"), cap)
    doc.save(OUT)

    # Грубая оценка объёма
    text = re.sub(r"```.*?```", "", md, flags=re.S)
    words = len(re.findall(r"\S+", text))
    print(f"Сохранено: {OUT}")
    print(f"Абзацев: {len(doc.paragraphs)}, таблиц: {len(doc.tables)}, "
          f"рисунков: {len(doc.inline_shapes)}")
    print(f"Слов (без кода): ~{words}; ориентировочно страниц: ~{round(words/270)}")
    for ch, name in (("## 1 ", "Глава 1"), ("## 2 ", "Глава 2"), ("## 3 ", "Глава 3")):
        a = md.find(ch)
        if a < 0:
            continue
        nxt = md.find("\n## ", a + 5)
        seg = md[a: nxt if nxt > 0 else len(md)]
        w = len(re.findall(r"\S+", re.sub(r"```.*?```", "", seg, flags=re.S)))
        print(f"  {name}: ~{w} слов (~{round(w/270)} стр.)")


if __name__ == "__main__":
    main()
