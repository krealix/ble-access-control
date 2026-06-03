"""Заполнение официальных форм ВКР (титул, задание, календарный план).

Открывает исходные формы из Downloads, вписывает известные поля (направление,
профиль, тема, год, исходные данные, содержание, план с датами), сохраняет
заполненные копии в diploma/. Личные поля (Ф.И.О., подписи) оставлены пустыми.
"""
import os

from docx import Document

DL = r"C:\Users\krealix\Downloads"
OUT = r"C:\Users\krealix\Desktop\ble\diploma"

NAPR = "01.04.02 Прикладная математика и информатика"
PROF = "Математическая робототехника и искусственный интеллект"
TEMA = ("Разработка системы контроля удалённого доступа на основе анализа "
        "траектории изменения сигнала BLE-меток")
YEAR = "2026"
SROK = "16.06.2026"

ISH = ("BLE-модуль HM10 (прозрачный мост BLE → RS-485); методические указания ЮГУ "
       "по выполнению ВКР; ГОСТ 34.602-2020; научные публикации по технологии BLE, "
       "оценке расстояния по RSSI, фильтрации Калмана и контролю доступа.")
SODER = ("введение; 1 Анализ предметной области и методов обработки сигнала "
         "BLE-меток; 2 Проектирование системы контроля доступа на основе анализа "
         "траектории сигнала; 3 Программная реализация и тестирование системы; "
         "заключение; список источников; приложения.")
GRAPH = ("структурная схема системы; граф состояний алгоритма; блок-схема алгоритма; "
         "графики результатов моделирования; презентация (не более 10 слайдов).")

PLAN = [
    ("Подбор и анализ литературы, постановка задачи", "10.02.2026"),
    ("Глава 1. Анализ предметной области и методов", "01.03.2026"),
    ("Глава 2. Проектирование системы и алгоритма", "25.03.2026"),
    ("Глава 3. Программная реализация системы", "25.04.2026"),
    ("Тестирование системы, обработка результатов", "12.05.2026"),
    ("Оформление пояснительной записки", "30.05.2026"),
    ("Подготовка презентации и доклада", "06.06.2026"),
    ("Предзащита ВКР", "10.06.2026"),
    ("Нормоконтроль", "13.06.2026"),
    ("Защита ВКР", "19.06.2026"),
]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def is_blank(t):
    s = t.strip()
    return s != "" and set(s) <= set("_ ")


def fill_dir_prof(paras):
    for i, p in enumerate(paras):
        if "код и наименование направления" in p.text and i > 0:
            set_text(paras[i - 1], NAPR)
        if "наименование профиля" in p.text and i > 0:
            set_text(paras[i - 1], PROF)


def clear_following_blanks(paras, i, limit=3):
    for j in range(i + 1, min(i + 1 + limit, len(paras))):
        if is_blank(paras[j].text):
            set_text(paras[j], "")


# ---------- Титульный лист ----------
def fill_titul():
    doc = Document(os.path.join(DL, "Форма Титульного листа_Магистратура.docx"))
    paras = doc.paragraphs
    fill_dir_prof(paras)
    for p in paras:
        if p.text.strip().startswith("На тему"):
            set_text(p, "На тему: " + TEMA)
        elif "20__год" in p.text.replace(" ", "") or p.text.strip() in ("20__год", "20__ год"):
            set_text(p, YEAR + " год")
    doc.save(os.path.join(OUT, "Титульный_лист.docx"))
    print("OK Титульный_лист.docx")


# ---------- Задание на ВКР ----------
def fill_zadanie():
    doc = Document(os.path.join(DL, "Форма задания на ВКР (магистратура).docx"))
    paras = doc.paragraphs
    fill_dir_prof(paras)
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == "1. Тема" and i + 1 < len(paras):
            set_text(paras[i + 1], "«" + TEMA + "»")
        elif "Срок сдачи студентом" in t:
            set_text(p, "2. Срок сдачи студентом законченной выпускной "
                        "квалификационной работы — " + SROK)
            clear_following_blanks(paras, i)
        elif "Исходные данные к выпускной" in t:
            set_text(p, "3. Исходные данные к выпускной квалификационной работе: " + ISH)
            clear_following_blanks(paras, i)
        elif t.startswith("4. Содержание"):
            set_text(p, "4. Содержание выпускной квалификационной работы (перечень "
                        "подлежащих разработке вопросов, разделов): " + SODER)
            clear_following_blanks(paras, i)
        elif t.startswith("5. Ориентировочный перечень графического"):
            set_text(p, "5. Ориентировочный перечень графического и иллюстративного "
                        "материала: " + GRAPH)
            clear_following_blanks(paras, i)
        elif t.startswith("7. Дата выдачи"):
            set_text(p, "7. Дата выдачи задания «10» февраля 2026 г.")
    doc.save(os.path.join(OUT, "Задание_на_ВКР.docx"))
    print("OK Задание_на_ВКР.docx")


# ---------- Календарный план ----------
def fill_plan():
    doc = Document(os.path.join(DL, "Форма календарного плана выполнения ВКР (для магистратуры).docx"))
    paras = doc.paragraphs
    fill_dir_prof(paras)
    for p in paras:
        if p.text.strip().startswith("Тема «"):
            set_text(p, "Тема «" + TEMA + "»")
    if doc.tables:
        t = doc.tables[0]
        body_rows = t.rows[1:]  # без шапки
        for row, (name, date) in zip(body_rows, PLAN):
            set_text(row.cells[0].paragraphs[0], name)
            set_text(row.cells[1].paragraphs[0], date)
    doc.save(os.path.join(OUT, "Календарный_план.docx"))
    print("OK Календарный_план.docx")


if __name__ == "__main__":
    fill_titul()
    fill_zadanie()
    fill_plan()
