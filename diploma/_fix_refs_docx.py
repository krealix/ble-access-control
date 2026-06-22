# -*- coding: utf-8 -*-
"""Перенумерация источников ВКР (docx) по порядку первого упоминания.
Убирает 5 нецитированных ГОСТов, переупорядочивает авто-нумерованный список,
переписывает внутритекстовые ссылки [N]. Сохраняет в новый файл."""
import re, sys
from docx import Document
sys.stdout.reconfigure(encoding="utf-8")

P_IN = r"D:\Claude projects\ВКР\Диплом — копия.docx"
P_OUT = r"D:\Claude projects\ВКР\Диплом (источники по порядку).docx"

# old -> new (по порядку первого появления в тексте)
MAP = {5: 1, 6: 2, 7: 3, 8: 4, 9: 5, 10: 6, 11: 7, 12: 8, 13: 9,
       14: 10, 15: 11, 16: 12, 17: 13, 18: 14, 1: 15}

CITE_PARAS = [92, 95, 97, 98, 102, 104, 113, 122, 125, 126, 127, 129, 130,
              132, 136, 137, 142, 146, 154, 166, 170, 174, 175, 181, 188,
              190, 330, 331]

cite_re = re.compile(r"\[(\d+(?:\s*[–\-,]\s*\d+)*)\]")


def map_inner(inner):
    out = []
    for part in inner.split(","):
        part = part.strip()
        m = re.match(r"^(\d+)\s*([–-])\s*(\d+)$", part)
        if m:
            a, dash, b = MAP[int(m.group(1))], m.group(2), MAP[int(m.group(3))]
            out.append(f"{a}{dash}{b}")
        else:
            out.append(str(MAP[int(part)]))
    return ", ".join(out)


def repl(m):
    return "[" + map_inner(m.group(1)) + "]"


d = Document(P_IN)
paras = d.paragraphs

# 1) переписать внутритекстовые ссылки (форматирование однородно — собираем в 1 run)
for i in CITE_PARAS:
    para = paras[i]
    new_text = cite_re.sub(repl, para.text)
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)

# 2) список: переместить flutter (452) после Sung (469); удалить 5 ГОСТов
flutter_p = paras[452]._p
sung_p = paras[469]._p
del_ps = [paras[k]._p for k in (453, 454, 455, 470, 471)]
sung_p.addnext(flutter_p)
for dp in del_ps:
    dp.getparent().remove(dp)

d.save(P_OUT)
print("Сохранено:", P_OUT)
